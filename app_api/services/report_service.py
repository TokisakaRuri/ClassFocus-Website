from __future__ import annotations

from pathlib import Path
from typing import Any

from app_api.core.config import settings
from app_api.services.statistics_service import ALL_CLASSES, BEHAVIOR_LABELS, derive_overall_evidence, derive_segment_evidence


BOUNDARY_TEXT = (
    "解释边界：系统输出仅代表可观察课堂行为及其时序证据，不测量学生内在认知专注状态；"
    "低头、书写、阅读、听课和使用手机等行为须结合教学任务由教师解释。结果不得用于个体排名、标签化评价或惩罚性决策。"
)


def export_report(
    task_id: int,
    course: dict[str, Any],
    video_info: dict[str, Any],
    overall: dict[str, Any],
    segments: list[dict[str, Any]],
    agent_report: dict[str, str],
    quality_report: dict[str, Any] | None = None,
    multi_agent: dict[str, Any] | None = None,
    artifact_key: str | None = None,
) -> Path:
    settings.report_dir.mkdir(parents=True, exist_ok=True)
    suffix = f"_{artifact_key}" if artifact_key else ""
    docx_path = settings.report_dir / f"classfocus_report_{task_id}{suffix}.docx"
    try:
        from docx import Document
    except ImportError:
        text_path = settings.report_dir / f"classfocus_report_{task_id}{suffix}.txt"
        return _export_text(text_path, course, video_info, overall, segments, agent_report)

    evidence = derive_overall_evidence(overall)
    document = Document()
    document.add_heading("ClassFocus 课堂行为证据分析报告", level=1)
    document.add_paragraph(BOUNDARY_TEXT)

    document.add_heading("一、课程信息", level=2)
    for label, key in (
        ("课程名称", "course_name"), ("教师姓名", "teacher_name"), ("班级名称", "class_name"),
        ("教室", "classroom"), ("上课日期", "lesson_date"), ("上课节次", "lesson_section"),
    ):
        document.add_paragraph(f"{label}：{course.get(key, '')}")

    document.add_heading("二、数据与模型信息", level=2)
    document.add_paragraph(f"视频时长：{video_info.get('duration', 0)} 秒")
    document.add_paragraph(f"帧率：{video_info.get('fps', 0)}")
    document.add_paragraph(f"分辨率：{video_info.get('resolution', '')}")
    document.add_paragraph(f"分析模式：{video_info.get('mode', '')}")

    document.add_heading("三、行为证据概览", level=2)
    document.add_paragraph(f"行为证据总数：{evidence['total_count']}")
    document.add_paragraph(f"行为类别覆盖：{evidence['class_count']}/{len(ALL_CLASSES)}")
    dominant = BEHAVIOR_LABELS.get(evidence["dominant_behavior"], "暂无")
    document.add_paragraph(f"主要行为：{dominant}（{evidence['dominant_behavior_rate']:.2f}%）")
    for class_name in ALL_CLASSES:
        document.add_paragraph(
            f"{BEHAVIOR_LABELS[class_name]}：{evidence['raw_counts'][class_name]} 次，"
            f"占比 {evidence['behavior_distribution'][class_name]:.2f}%"
        )

    if quality_report:
        document.add_heading("四、证据完整性", level=2)
        document.add_paragraph(quality_report.get("summary", ""))
        for dimension in quality_report.get("dimension_scores", []):
            document.add_paragraph(
                f"{dimension.get('name', '')}：{dimension.get('score', 0)}%。{dimension.get('evidence', '')}"
            )

    document.add_heading("五、时间段证据", level=2)
    table = document.add_table(rows=1, cols=6)
    headers = ["时间段", "检测数", "主要行为", "主要行为占比", "复核优先级", "复核依据"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for segment in segments[:80]:
        item = derive_segment_evidence(segment)
        cells = table.add_row().cells
        cells[0].text = str(segment.get("time_range", ""))
        cells[1].text = str(item["total_count"])
        cells[2].text = BEHAVIOR_LABELS.get(item["dominant_behavior"], "暂无")
        cells[3].text = f"{item['dominant_behavior_rate']:.2f}%"
        cells[4].text = item["review_priority"]
        cells[5].text = item["review_reason"]

    document.add_heading("六、三方面辅助分析", level=2)
    document.add_paragraph(agent_report.get("full_report", ""))
    if multi_agent:
        document.add_heading("七、协同分析结论", level=2)
        document.add_paragraph(multi_agent.get("consensus", ""))
    document.save(docx_path)
    return docx_path


def _export_text(
    path: Path,
    course: dict[str, Any],
    video_info: dict[str, Any],
    overall: dict[str, Any],
    segments: list[dict[str, Any]],
    agent_report: dict[str, str],
) -> Path:
    evidence = derive_overall_evidence(overall)
    dominant = BEHAVIOR_LABELS.get(evidence["dominant_behavior"], "暂无")
    lines = [
        "ClassFocus 课堂行为证据分析报告", "", BOUNDARY_TEXT, "",
        f"课程名称：{course.get('course_name', '')}",
        f"班级名称：{course.get('class_name', '')}",
        f"视频时长：{video_info.get('duration', 0)} 秒",
        f"行为证据总数：{evidence['total_count']}",
        f"行为类别覆盖：{evidence['class_count']}/{len(ALL_CLASSES)}",
        f"主要行为：{dominant}（{evidence['dominant_behavior_rate']:.2f}%）", "",
        agent_report.get("full_report", ""), "", "时间段证据",
    ]
    for segment in segments:
        item = derive_segment_evidence(segment)
        lines.append(
            f"{segment.get('time_range', '')} | 检测数 {item['total_count']} | "
            f"主要行为 {BEHAVIOR_LABELS.get(item['dominant_behavior'], '暂无')} | "
            f"复核优先级 {item['review_priority']} | {item['review_reason']}"
        )
    path.write_text("\n".join(lines), encoding="utf-8")
    return path
